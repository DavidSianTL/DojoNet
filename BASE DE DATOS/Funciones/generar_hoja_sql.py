from docx import Document

# Crear documento
doc = Document()
doc.add_heading('Hoja de Ejercicios – Funciones en SQL Server', 0)

# Introducción
doc.add_paragraph(
    "Esta hoja de ejercicios tiene como objetivo reforzar el uso de diferentes tipos de funciones en SQL Server 2019, "
    "incluyendo funciones escalares, agregadas, de texto, de fecha, de ventana y definidas por el usuario.\n"
)

# Sección 1: Funciones Matemáticas Escalares
doc.add_heading('1. Funciones Matemáticas Escalares', level=1)
doc.add_paragraph("Usa la tabla Productos ya creada. Escribe las consultas para:")
doc.add_paragraph("a) Mostrar el valor absoluto del precio de cada producto.\n"
                  "b) Redondear el precio hacia arriba y hacia abajo.\n"
                  "c) Calcular la raíz cuadrada y el cuadrado del precio.\n"
                  "d) Obtener el signo del precio y un número aleatorio por producto.")

# Sección 2: Funciones de Texto
doc.add_heading('2. Funciones de Texto', level=1)
doc.add_paragraph("Trabajando con la columna Nombre de la tabla Productos:")
doc.add_paragraph("a) Convertir todos los nombres a mayúsculas.\n"
                  "b) Obtener la longitud del nombre del producto.\n"
                  "c) Extraer las primeras 5 letras del nombre del producto.\n"
                  "d) Reemplazar la palabra 'Samsung' por 'LG'.")

# Sección 3: Funciones de Fecha
doc.add_heading('3. Funciones de Fecha', level=1)
doc.add_paragraph("Agrega una columna FechaIngreso a la tabla Empleados. Luego:")
doc.add_paragraph("a) Mostrar la fecha actual del sistema.\n"
                  "b) Calcular cuántos días lleva cada empleado en la empresa.\n"
                  "c) Sumar 30 días a la fecha de ingreso.\n"
                  "d) Mostrar solo el año de ingreso.")

# Sección 4: Funciones Agregadas
doc.add_heading('4. Funciones Agregadas', level=1)
doc.add_paragraph("Usando la tabla Productos:")
doc.add_paragraph("a) Calcular el precio promedio de todos los productos.\n"
                  "b) Mostrar el precio mínimo y máximo por categoría.\n"
                  "c) Contar cuántos productos hay por categoría.\n"
                  "d) Sumar el total de precios por categoría.")

# Sección 5: Funciones de Ventana
doc.add_heading('5. Funciones de Ventana', level=1)
doc.add_paragraph("Utilizando RANK, DENSE_RANK y ROW_NUMBER:")
doc.add_paragraph("a) Obtener un ranking global por precio (descendente).\n"
                  "b) Obtener un ranking por categoría usando RANK().\n"
                  "c) Mostrar el número de fila por categoría con ROW_NUMBER().\n"
                  "d) Comparar los resultados de RANK vs DENSE_RANK.")

# Sección 6: Funciones Definidas por el Usuario (UDF)
doc.add_heading('6. Funciones Definidas por el Usuario (UDF)', level=1)
doc.add_paragraph("Crea una función que reciba un salario mensual y devuelva el salario anual (13 pagos). Luego:")
doc.add_paragraph("a) Aplica la función a la tabla Empleados.\n"
                  "b) Modifica la función para aceptar un número de pagos como parámetro.\n"
                  "c) Usa la función en una consulta que calcule el salario anual con 14 pagos.")

# Guardar el documento
doc.save("Hoja_Ejercicios_Funciones_SQLServer.docx")
